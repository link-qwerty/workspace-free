# Import
import os as shell
import argparse as cmd
import smtplib as smtp_handler
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import docx
from datetime import datetime

# Defines
class MailerStreamer:
    """
    Класс для работы с почтовыми сообщениями

    Объекты класса формируют и отправляют почтовые сообщения в конвейере, используя протокол smtp. Входные данные
    представлены в виде файлов .docx, в которых содержится: служебный заголовок с адресами рассылки и темой сообщения,
    тело письма с изображением, подвал с отметками об отправке. Скрипт считывает файл, парсит информацию в нем и
    формирует почтовое сообщение, которое затем будет отправлено на адрес получателя, указанный в рассылке.

    """

    def __init__(self, host:str, port:int, work_directory:str, out_directory:str, tmp_directory:str, starttls:bool = False):
        self.host = host
        self.port = port
        self.work_directory = work_directory
        self.out_directory = out_directory
        self.tmp_directory = tmp_directory
        self.starttls = starttls

    def send_messages(self, sender_email:str, sender_password:str):
        # Собираем генератор списка для обхода файлов .docx в рабочей директории
        files = [f for f in shell.listdir(self.work_directory) if f.endswith('.docx')]

        # Перебираем файлы
        for docx_file in files:
            # Создаем объект документа и подгружаем в него файл
            docx_document = docx.Document(shell.path.join(self.work_directory, docx_file))

            # Создаем объект сообщения
            message = MIMEMultipart()
            # Результат отправки
            result = str()

            # Перебираем таблицы в документе
            for table in docx_document.tables:
                # Перебираем ячейки в таблице
                if table.cell(0,0).text == "SYSHEADER":

                    # Наполняем сообщение содержимым
                    message['From'] = sender_email
                    message['To'] = table.cell(1,1).text
                    message['Subject'] = table.cell(2,1).text
                    # Создаем тело сообщения
                    message_body = str()
                    # Записываем текст сообщения в тело
                    for paragraph in docx_document.paragraphs:
                        message_body += paragraph.text
                    message.attach(MIMEText(str(message_body), 'plain', 'utf-8'))
                    # Заходим на почтовый сервер
                    with smtp_handler.SMTP_SSL(self.host, self.port) as server:
                        if self.starttls:
                            server.starttls()
                        server.login(args.sender_email, args.sender_password)
                        # Отправляем сообщение
                        result = f'From: {message['From']}\nTo: {message['To']}\nDate:{datetime.now()}'
                        server.sendmail(message['From'], message['To'].split(','), message.as_string())
                        # Разлогиниваемся
                        server.quit()
                elif table.cell(0,0).text == "SYSFOOTER":
                    # Записываем информацию об отправке в документ
                    table.cell(1, 1).text = result
                    # Сохраняем документ в директории отработанных писем
                    docx_document.save(shell.path.join(self.out_directory, f'{datetime.now().date()} {datetime.now().timestamp()} {docx_file}'))

# Code

# Определение аргументов командной строки
parser = cmd.ArgumentParser(description="Mailer Queue(cl)LQ@2025 Free to use")
parser.add_argument("-s", "--server", help="SMTP server address",
                    default="smtp.mail.ru", type=str)
parser.add_argument("-P", "--port", help="SMTP server port",
                    default=465, type=int)
parser.add_argument("-e", "--sender-email", help="sender's e-mail",
                    required=True, type=str)
parser.add_argument("-p", "--sender-password", help="sender's password",
                    required=True, type=str)
parser.add_argument("-w", "--work-directory",help="directory for task files",
                    default=shell.path.abspath("work/"), type=str)
parser.add_argument("-o", "--out-directory", help="directory for completed tasks",
                    default=shell.path.abspath("out/"), type=str)
parser.add_argument("-l", "--logs-directory", help="directory for logs",
                    default=shell.path.abspath("logs/"), type=str)
parser.add_argument("-t", "--tmp-directory", help="directory for temp files",
                    default=shell.path.abspath("tmp/"), type=str)

# Парсинг аргументов
args = parser.parse_args()
# Создание рабочих директорий (если не существуют)
shell.makedirs(args.work_directory, exist_ok = True)
shell.makedirs(args.out_directory, exist_ok = True)
shell.makedirs(args.logs_directory, exist_ok = True)
shell.makedirs(args.tmp_directory, exist_ok = True)
# Создание объекта-мейлера
mailer = MailerStreamer(args.server, args.port, args.work_directory, args.out_directory, args.tmp_directory)
# Отправка сообщений из директории
mailer.send_messages(args.sender_email, args.sender_password)